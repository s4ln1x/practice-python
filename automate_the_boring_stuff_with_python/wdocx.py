#! /usr/bin/env python3
import docx

# Read a docx file
d = docx.Document("Tesis.docx")

# d.paragraphs has all the paragraphs objects of the document
for paragraph in d.paragraphs:
    print(paragraph.text)

# Each paragraph has an object named runs
p = d.paragraphs[0]

# Paragraphs have a variable named style that stores the style of the text
p.style

for run in p.runs:
    print(run.text)

# This variables inside the run object are booleans and hold the information of
# the format
p.runs[0].bold
p.runs[0].italic
p.runs[0].underline
p.runs[0].bold

# save method saves modifications to the document
p.runs[0].text = "Hell yeah pops I aded this xD"
p.runs[0].bold = True
p.runs[0].italic = True
d.save("Tesis.docx")

# Write my own docx file
d_2 = docx.Document()

# Add new paragraphs
d_2.add_paragraph("Hey pops what are you doing with your money?")
d_2.add_paragraph(
    "You can invest your money in Forex or maybe buy some bitcoin xD"
)

# Add new runs to the first paragraph
p_2 = d_2.paragraphs[0]

p_2.add_run("$Bitcoin$ = HELL YEAH!!!")
p_2.runs[1].underline = True

d_2.save("Bitcoins.docx")
